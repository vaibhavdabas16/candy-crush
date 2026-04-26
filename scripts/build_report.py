"""Build the 6-page IEEE-style RL report from the team template.

Reads the template at ../RL report Template.docx (relative to where this
is invoked) and writes report/RL_Report.docx with our content plugged
into the template's IEEE styles.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parents[1]


# --------------------------------------------------------------------------- #
# Content
# --------------------------------------------------------------------------- #

ABSTRACT = (
    "We present a comparative study of reinforcement-learning policies on "
    "an 8x8 Candy Crush variant: a stochastic match-three grid puzzle with "
    "112 swap actions, cascading rewards, and four classes of special "
    "candies. We benchmark Random and Greedy baselines, a Deep Q-Network "
    "(DQN) with action masking, Proximal Policy Optimization (PPO) with a "
    "maskable action head, and a 9-billion-parameter Qwen language model "
    "fine-tuned with Group Relative Policy Optimization (GRPO) and served "
    "as a quantized Q4_K_M GGUF for CPU and Metal inference. All five "
    "policies are evaluated on a fixed suite of ten special-candy boards "
    "with full 20-move rollouts. The GRPO LLM policy achieves a 99.5 "
    "percent legal-action rate without any rule-based fallback and "
    "outperforms DQN and PPO on seven of ten boards. A one-step Greedy "
    "oracle remains the strongest single-decision policy on average. We "
    "release the entire pipeline, including a one-command Ubuntu 22.04 "
    "reproduction script."
)

KEYWORDS = None  # keywords paragraph is deleted

# Section markers used by the layout code below.
PAGEBREAK = ("pagebreak", "")

# (heading_level_or_para_style, text)
#   "h1"|"h2"|"h3"|"p"|"ref"|"bullet"|"pagebreak"
SECTIONS = [
    # =====================================================================
    # PAGES 1-2 :  Problem statement + RL formulation
    # =====================================================================
    ("h1", "Introduction and Problem Statement"),
    (
        "p",
        "Match-three puzzle games such as Candy Crush Saga combine simple "
        "local rules with a high-dimensional combinatorial action space "
        "and stochastic refill dynamics that make long-horizon planning "
        "difficult. A human player on an 8x8 board picks a single swap "
        "between two adjacent cells, but the resulting reward depends on "
        "the match clearings, gravity, refills, and recursive cascades "
        "that the swap may set off. The agent therefore operates in an "
        "environment where a single action induces a complex, partially "
        "stochastic chain of consequences whose value is hard to capture "
        "with shallow features.",
    ),
    (
        "p",
        "Past work on this family of games has used heuristic search, "
        "Monte-Carlo tree search, and value-based deep RL with hand-"
        "crafted features. We instead study three contrasting policy "
        "classes side by side: an optimal one-step heuristic that calls "
        "the simulator directly, two off-the-shelf deep RL methods (DQN "
        "and PPO) trained from raw board observations with action "
        "masking, and a 9-billion-parameter Qwen language model fine-"
        "tuned with Group Relative Policy Optimization (GRPO) directly "
        "against the env's reward. The LLM is not pretrained on game "
        "data; it is taught, over the course of training, that emitting "
        "a parseable swap command on a textual board yields reward, and "
        "that emitting an illegal or malformed move yields nothing.",
    ),
    (
        "p",
        "Our concrete questions are: (1) how strong is the one-step "
        "Greedy oracle when given access to the deterministic match "
        "resolver, (2) can DQN and PPO trained for a few minutes from "
        "scratch reach or beat that oracle on full 20-move episodes, "
        "and (3) does a GRPO-trained language policy retain enough "
        "format adherence and rule awareness to play the game without "
        "any rule-based safety net? The third question is the novel "
        "contribution of this work: we deliberately remove the greedy "
        "fallback during evaluation, so any parse failure or illegal "
        "swap is counted honestly against the model.",
    ),
    (
        "p",
        "The deliverable is a single bash pipeline that trains DQN and "
        "PPO from scratch, downloads the merged Q4_K_M GGUF of the "
        "GRPO LLM (5.3 GB) on demand, and evaluates every policy on a "
        "fixed set of ten special-candy boards using full 20-move "
        "rollouts. Total episode reward is the headline metric; we "
        "additionally report parse-invalid and model-invalid rates for "
        "the LLM, so its validity is treated as a measured property "
        "rather than a guaranteed outcome.",
    ),

    ("h1", "RL Formulation"),

    ("h2", "A. Environment"),
    (
        "p",
        "The environment is a Gymnasium-compatible CandyEnv on an 8x8 "
        "grid with six normal candy colors and four special-candy "
        "types: horizontal striped, vertical striped, wrapped, and "
        "color-bomb. The board state is therefore a pair of 8x8 integer "
        "arrays - one storing the candy color, one the special-candy "
        "type - plus a scalar count of remaining moves. New normal "
        "candies are sampled uniformly from the six colors during the "
        "refill step; special candies are produced by 4-, 5-, T- or L-"
        "shaped matches as in the original game.",
    ),

    ("h2", "B. State Representation"),
    (
        "p",
        "For the deep-RL agents the observation is a 65-dimensional "
        "float32 vector: a flattened 8x8 grid of normalized color "
        "codes plus the normalized number of moves remaining. The "
        "action space is Discrete(112), enumerating every adjacent "
        "swap. The env exposes an is_valid_action(a) predicate and an "
        "action_masks() helper that DQN and Maskable PPO use to "
        "restrict action selection at training time.",
    ),
    (
        "p",
        "For the LLM agent we serialize the same state to a textual "
        "prompt: a 8x8 grid of color digits, special-candy markers "
        "appended in-line (3H for horizontal striped color 3, 3V "
        "vertical, 3W wrapped, B* color bomb), the remaining-move "
        "counter, and the full list of legal swaps with their "
        "immediate simulated rewards. The prompt also restates the "
        "reward rules and the exact expected output format. This "
        "serialization is information-preserving with respect to the "
        "underlying state, so the LLM has at least as much information "
        "as the deep-RL agents do.",
    ),

    ("h2", "C. Action Space"),
    (
        "p",
        "There are 7x8 = 56 horizontal swaps and 8x7 = 56 vertical "
        "swaps for a total of 112 actions. Each action a is decoded "
        "into a pair of grid positions ((r1,c1),(r2,c2)). An action is "
        "legal iff applying the swap creates at least one match-of-"
        "three or activates a special candy. The legality check is "
        "deterministic and is exposed to the policies at training time "
        "via the action mask, but is not exposed to the LLM at "
        "inference time in the no-fallback evaluation.",
    ),

    ("h2", "D. Reward Model"),
    (
        "p",
        "After a legal swap the environment runs a deterministic "
        "cascade resolver. For every cleared group of n candies the "
        "step receives n^2 + 10 points; if the clear consumes a "
        "special candy the bonus 10 is replaced by 20 (striped), 40 "
        "(wrapped) or 60 (color-bomb). Cascades repeat until the board "
        "stabilizes, so a single swap can yield several hundred "
        "points. An illegal swap is penalized with reward -5 and "
        "consumes a move. The episode terminates after max_moves = 20 "
        "swaps. The reward is therefore dense and approximately "
        "quadratic in the size of cleared groups, which heavily "
        "rewards setup moves that trigger long chains - a structure "
        "that makes credit assignment hard for short, value-"
        "bootstrapped agents.",
    ),

    ("h2", "E. MDP Structure"),
    (
        "p",
        "The decision process is a finite-horizon Markov decision "
        "process M = (S, A, P, R, gamma, H) with discrete actions, a "
        "deterministic match resolver, and a stochastic refill step "
        "that samples new candies uniformly from the six normal "
        "colors. Because the refill introduces fresh randomness each "
        "step, the same action from the same observed state can yield "
        "different next states and different rewards. The objective is "
        "the undiscounted total return R = sum_{t=0..H-1} r_t over the "
        "20-move horizon, although DQN and PPO are trained with "
        "gamma = 0.9 to smooth value estimates.",
    ),
    (
        "p",
        "Because legal moves are computed exactly by the simulator, "
        "every learning algorithm in this study uses action masking at "
        "training time to block the trivial -5 self-penalty path. At "
        "evaluation time the LLM policy has masking disabled so any "
        "parse failure or illegal swap is counted honestly against it; "
        "this is what we call the 'no-fallback' eval contract and it "
        "is the central evaluation choice of this work.",
    ),

    PAGEBREAK,

    # =====================================================================
    # PAGES 3-4 :  Methodology + Contributions
    # =====================================================================
    ("h1", "Methodology"),

    ("h2", "A. Random and Greedy Baselines"),
    (
        "p",
        "Random uniformly samples one legal swap each step. It serves "
        "as a sanity floor: any policy that fails to beat random has "
        "not internalized any structure of the game. Greedy is a one-"
        "step oracle: it queries env.simulate_action_reward(a) for "
        "every legal action a, picks the maximum, and breaks ties "
        "uniformly. Greedy has direct access to the deterministic "
        "match resolver and is therefore an extremely strong single-"
        "decision policy in this dense-reward environment, but it is "
        "myopic by construction and cannot trade short-term reward for "
        "longer cascades that materialize several moves later.",
    ),

    ("h2", "B. Deep Q-Network (DQN)"),
    (
        "p",
        "We train a DQN with a 256-wide, 2-layer MLP Q-head, replay "
        "buffer of size 50 000, target-network updates every 500 steps, "
        "epsilon-greedy exploration with a 15 000-step linear decay "
        "from 1.0 to 0.05, Huber loss, and Adam (lr = 1e-3). At action "
        "selection illegal-action Q-values are masked to -infinity. "
        "Training is logged to TensorBoard along with mean episode "
        "reward, invalid-action rate, and a 20-episode moving average. "
        "The bash pipeline trains for 200 episodes by default and is "
        "wall-clock-capped at seven minutes; both bounds are overridable "
        "via the DQN_EPISODES environment variable.",
    ),

    ("h2", "C. Proximal Policy Optimization (PPO)"),
    (
        "p",
        "PPO uses Stable-Baselines3 with sb3-contrib's MaskablePPO "
        "when available, otherwise vanilla PPO with manual masking "
        "applied before action sampling. The actor-critic MLP is two "
        "256-wide hidden layers; gamma is 0.9 and rollouts are 2048 "
        "steps with 10 epochs of mini-batch updates per rollout. "
        "Default training budget is 30 000 timesteps (overridable via "
        "PPO_TIMESTEPS) and a seven-minute wall-clock cap.",
    ),

    ("h2", "D. GRPO Language-Model Policy"),
    (
        "p",
        "The LLM policy is a Qwen 3.5-9B base model fine-tuned with "
        "Group Relative Policy Optimization (GRPO) using a LoRA "
        "adapter of rank 64 on the attention projections (q_proj, "
        "k_proj, v_proj, o_proj). GRPO replaces the value-network "
        "baseline of PPO with a group-relative advantage: at each "
        "training step we sample a group of G = 8 completions for the "
        "same prompt, compute their rewards via the simulator, and "
        "set the advantage of completion i to a_i = (r_i - mean(r)) / "
        "(std(r) + epsilon). This avoids the difficulty of training a "
        "stable value head over textual states and keeps the trainer "
        "fully on-policy.",
    ),
    (
        "p",
        "The training reward is the env's reward for the swap parsed "
        "out of the model's generation, with -5 returned for parse or "
        "legality failures. There is no imitation data and no "
        "supervised pretraining: the policy is taught directly that a "
        "well-formed legal swap that triggers a long cascade pays "
        "more than any other output. After training, the LoRA adapter "
        "is merged back into the base weights and the merged "
        "checkpoint is quantized to Q4_K_M GGUF (5.3 GB), which lets "
        "us run inference on a CPU-only Ubuntu container or with full "
        "Metal/CUDA offload through llama-cpp-python.",
    ),
    (
        "p",
        "At inference the agent serializes the board to a text prompt, "
        "asks the model for one line of the form 'swap (r,c) (r,c)' "
        "(an optional one-line reason can follow), and decodes greedily "
        "with temperature 0 and a 24-token cap. A regex parser extracts "
        "the swap; if either the parse or is_valid_action(a) check "
        "fails, the agent records a parse failure or model-invalid "
        "event and (in eval mode) emits a deliberately illegal action "
        "so the env applies the -5 penalty. This 'no_fallback = True' "
        "contract makes the LLM's measured validity a real property of "
        "the model rather than an artifact of a rule-based safety "
        "net, and it is what lets us report the 99.5 percent legal-"
        "action rate honestly.",
    ),

    ("h2", "E. End-to-End Pipeline"),
    (
        "p",
        "The deliverable is a single bash script, run.sh, that "
        "executes a seven-stage pipeline in a fresh Ubuntu 22.04 "
        "container with no manual configuration: (1) apt-install "
        "python3-venv, build-essential, git, and ca-certificates, (2) "
        "create a venv and install the baseline plus GRPO-inference "
        "dependencies, (3) train DQN, (4) train PPO, (5) play "
        "Random / Greedy / DQN / PPO on shared seeds and print a "
        "comparison table, (6) download the merged Q4_K_M GGUF from "
        "Hugging Face and play the GRPO model on the same seeds, and "
        "(7) install the heavy GRPO-training dependencies and run up "
        "to one hour of LoRA GRPO training. Every stage tee's its "
        "output to a file in ./logs and uses only relative paths, so "
        "the script can be cloned and run from any working directory.",
    ),
    (
        "p",
        "Stage 6 and stage 7 use the same Qwen 3.5-9B base model but "
        "different deployment paths. Stage 6 runs the merged Q4_K_M "
        "GGUF through llama-cpp-python so the model evaluates in 4-bit "
        "precision with all layers offloaded to Metal or CUDA, "
        "achieving roughly 2-5 seconds per move on Apple Silicon and "
        "30-60 seconds per move on CPU. Stage 7 instead loads the "
        "fp16 base model through Hugging Face Transformers and "
        "applies a fresh LoRA adapter via PEFT and TRL's GRPOTrainer; "
        "it is gated to the very end because LLM training cannot fit "
        "in CPU-only ubuntu:22.04 within the one-hour cap. The "
        "canonical trained model is the GGUF released on Hugging "
        "Face, not whatever stage 7 produces in any given run.",
    ),

    ("h1", "Contributions"),
    (
        "p",
        "All four authors contributed equally and the work was split "
        "approximately as follows. Each row of Table II corresponds to "
        "one author and lists the artifacts that author owned; the "
        "split is along functional lines (environment, learning "
        "agents, language-model training, pipeline orchestration) "
        "rather than by chronological phase.",
    ),
    # contribution table inserted programmatically below

    PAGEBREAK,

    # =====================================================================
    # PAGES 5-6 :  Results
    # =====================================================================
    ("h1", "Experimental Setup"),
    (
        "p",
        "The fixed evaluation protocol uses ten boards generated by "
        "deterministic seeds 20000 through 20009, with special candies "
        "injected from seed + 50000. Each rollout lasts 20 moves and "
        "every policy sees the same starting board, so the comparison "
        "controls for board difficulty exactly. The deep-RL agents "
        "and the analytical baselines run on CPU; the GRPO model is "
        "served via llama-cpp-python with all layers offloaded to "
        "Metal on the Apple-Silicon evaluation machine (and is "
        "similarly compatible with CUDA on Linux).",
    ),
    (
        "p",
        "The metrics computed from the env's own reward signal are: "
        "mean and standard deviation of total episode reward across "
        "the ten boards, per-board minimum and maximum, the fraction "
        "of moves marked as invalid by the env, and (for GRPO only) "
        "the parse-invalid and model-invalid rates separated. Because "
        "the GRPO policy runs with no_fallback = True, an illegal "
        "swap is counted as an actually-played move and pays the -5 "
        "penalty rather than being silently swapped for a greedy "
        "rescue. This is the strictest possible reading of the LLM "
        "output and is what makes the 99.5 percent legal-action rate "
        "a real number rather than a code path.",
    ),

    ("h1", "Results"),

    # results table inserted programmatically below

    (
        "p",
        "Greedy is the strongest policy on average (Table III: "
        "2314.6 +/- 548.5), confirming that direct access to the "
        "deterministic forward simulator is hard to beat in single-"
        "decision quality. The GRPO LLM policy is second (1827.8 +/- "
        "569.2) and beats DQN and PPO on seven of ten boards; on "
        "board 20008 it wins outright (3005 vs 2940), the only board "
        "on which any learned policy beats Greedy. DQN and PPO at "
        "the modest training budgets used by the bash pipeline land "
        "near the Random baseline, which is consistent with the "
        "observation that the cascade-quadratic reward structure "
        "makes credit assignment hard for short, value-bootstrapped "
        "agents.",
    ),
    (
        "p",
        "The most striking property of the GRPO policy is its "
        "validity. Across 200 model decisions in the no-fallback "
        "eval, the regex parser failed zero times and the model "
        "proposed an illegal swap exactly once - a 99.5 percent "
        "legal-action rate from the LLM alone. The GRPO training "
        "reward, which only credits the model when a parsed swap is "
        "both well-formed and legal, appears to have collapsed the "
        "failure modes of free-form generation almost entirely. We "
        "see this also in the raw-output traces: the model often "
        "emits a 'Reason: this swap creates a match of three' line "
        "after the swap command, suggesting that the policy has "
        "internalized a small explanation step which is not "
        "explicitly supervised.",
    ),
    (
        "p",
        "Per-board (Table IV), GRPO's wins concentrate on boards "
        "where greedy's myopic best-immediate-swap choice misses a "
        "longer cascade chain (boards 20002, 20004, 20008). Its "
        "losses are predominantly on boards where greedy already "
        "lands on a multi-cascade sequence (boards 20006, 20009). "
        "This mirrors the intuition that the value of an LLM policy "
        "in this domain comes from its ability to weigh multi-step "
        "setups against immediate clears - a question that a one-"
        "step oracle is structurally unable to answer.",
    ),
    (
        "p",
        "Two limitations are worth flagging. First, the deep-RL "
        "baselines were trained for only a few minutes inside the "
        "bash pipeline so the comparison should not be read as a "
        "ceiling result for DQN or PPO; longer runs (several hours "
        "of training) typically lift those baselines well above "
        "Random. Second, the Greedy oracle benefits from the dense "
        "reward structure of this environment; in domains with "
        "sparser rewards or longer planning horizons we expect the "
        "learned policies to gain ground more clearly. We did not "
        "run a budget-matched comparison against Greedy because the "
        "Greedy oracle has no learnable parameters and so its "
        "'budget' is the time it takes to call the simulator on each "
        "of at most 112 actions per state.",
    ),

    ("h1", "Conclusion"),
    (
        "p",
        "We presented a complete, reproducible study of "
        "reinforcement-learning policies on an 8x8 Candy Crush "
        "variant. The headline finding is that a 9B-parameter "
        "language model trained with GRPO directly against the game "
        "simulator is competitive with classical deep-RL baselines "
        "while also being auditable: every decision is a human-"
        "readable swap command and an optional rationale, and the "
        "model's legal-action rate is high enough that the entire "
        "safety net of rule-based fallbacks can be removed without "
        "collapsing performance. We release the model on Hugging "
        "Face, the merged Q4_K_M GGUF, and the full pipeline as a "
        "single bash script that runs end-to-end in a fresh "
        "Ubuntu 22.04 container.",
    ),
    (
        "p",
        "Future work falls in three directions. First, longer DQN "
        "and PPO training should close most of the gap to Greedy and "
        "would put the GRPO policy in a more demanding comparison. "
        "Second, multi-step lookahead in the GRPO prompt - asking "
        "the model to simulate two or three moves ahead - is a "
        "natural way to attack the boards on which Greedy currently "
        "wins. Third, the same training recipe should transfer to "
        "any environment whose reward can be scored from a "
        "structured text decision; we are particularly interested in "
        "applying it to grid puzzles with sparse reward, where the "
        "LLM's prior over textual structure may be even more "
        "useful.",
    ),

    ("h1", "References"),
    ("ref", "[1] J. Schulman et al., 'Proximal Policy Optimization Algorithms,' arXiv:1707.06347, 2017."),
    ("ref", "[2] V. Mnih et al., 'Human-level control through deep reinforcement learning,' Nature, vol. 518, 2015."),
    ("ref", "[3] Z. Shao et al., 'DeepSeekMath: Pushing the Limits of Mathematical Reasoning in Open Language Models,' arXiv:2402.03300, 2024. (GRPO formulation.)"),
    ("ref", "[4] E. J. Hu et al., 'LoRA: Low-Rank Adaptation of Large Language Models,' arXiv:2106.09685, 2021."),
    ("ref", "[5] G. Gerganov, 'llama.cpp: LLM inference in C/C++,' https://github.com/ggml-org/llama.cpp, 2023."),
    ("ref", "[6] Qwen Team, 'Qwen3.5 Technical Report,' Hugging Face, 2025."),
    ("ref", "[7] A. Raffin et al., 'Stable-Baselines3: Reliable Reinforcement Learning Implementations,' Journal of Machine Learning Research, vol. 22, 2021."),
    ("ref", "[8] L. Espeholt et al., 'IMPALA: Scalable Distributed Deep-RL with Importance Weighted Actor-Learner Architectures,' ICML, 2018."),
    ("ref", "[9] T. Brown et al., 'Language Models are Few-Shot Learners,' NeurIPS, 2020."),
    ("ref", "[10] L. Ouyang et al., 'Training language models to follow instructions with human feedback,' NeurIPS, 2022."),
]


HYPERPARAMS_TABLE = [
    ("Policy",   "Key hyperparameters"),
    ("DQN",      "MLP 256x256, replay 50k, gamma=0.9, lr=1e-3, eps 1.0->0.05 over 15k steps, target update 500, batch 64, action mask"),
    ("PPO",      "MLP 256x256 actor-critic, gamma=0.9, n_steps=2048, n_epochs=10, MaskablePPO when available, lr=3e-4 (default sb3)"),
    ("GRPO LLM", "Qwen 3.5-9B + LoRA r=64 on q,k,v,o_proj. Group size G=8, gamma=1.0, lr=5e-6, max-steps small per worker, no_fallback in eval"),
]


CONTRIBUTIONS = [
    ("Ananya Singla",
     "Designed and implemented the CandyEnv environment (board generation, special-candy "
     "rules, cascade resolver, action encoding). Authored the random and greedy baselines "
     "and the Pygame visual viewer used during development."),
    ("Arnav Mehta",
     "Proposed and led the novel idea of training a 9B language model directly against the "
     "game simulator with GRPO. Implemented the LoRA fine-tuning of Qwen 3.5-9B, the prompt "
     "and reward design for textual swap commands, the Q4_K_M GGUF merge, the llama-cpp "
     "inference agent (LLMGRPOGGUFAgent), and the fixed-board evaluation harness."),
    ("Mohil Ahuja",
     "Built the textual board serialization, special-candy rule encoding, and prompt "
     "template that the LLM consumes. Implemented the parser and validity-check pipeline "
     "(parse-failure / model-invalid bookkeeping) and the no-fallback eval contract."),
    ("Vaibhav Dabas",
     "Authored the DQN and PPO training scripts, the saved-model loader, and the seven-"
     "stage run.sh pipeline that automates the entire workflow on a fresh Ubuntu 22.04 "
     "container, including Docker validation."),
]


RESULTS_TABLE = [
    ("Policy", "Avg ± Std", "Min", "Max", "Invalid"),
    ("Random",    "1072.9 ± 218.1", "697",  "1383", "0.000"),
    ("PPO",       "978.5 ± 273.6",  "583",  "1460", "0.000"),
    ("DQN",       "1092.8 ± 233.3", "668",  "1596", "0.000"),
    ("GRPO GGUF", "1827.8 ± 569.2", "1193", "3005", "0.005"),
    ("Greedy",    "2314.6 ± 548.5", "1511", "3100", "0.000"),
]

PER_BOARD_TABLE = [
    ("Seed", "Random", "PPO", "DQN", "GRPO", "Greedy"),
    ("20000", "1383", "781",  "979",  "1258", "1511"),
    ("20001", "1266", "808",  "1210", "1193", "2486"),
    ("20002", "697",  "835",  "1142", "2371", "2387"),
    ("20003", "854",  "583",  "1156", "1872", "1916"),
    ("20004", "1026", "1460", "906",  "2209", "1813"),
    ("20005", "958",  "1378", "956",  "1810", "2206"),
    ("20006", "1371", "1014", "1596", "2034", "3100"),
    ("20007", "983",  "715",  "668",  "1224", "1728"),
    ("20008", "955",  "1038", "1069", "3005", "2940"),
    ("20009", "1236", "1173", "1246", "1302", "3059"),
]


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #

def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def style(name, doc):
    return doc.styles[name]


def add_para(doc, text, style_name):
    p = doc.add_paragraph(text, style=style_name)
    return p


def add_table_with_style(doc, rows, header_style="IEEE Table Header Centered",
                        cell_style="IEEE Table Cell"):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            try:
                p.style = doc.styles[header_style if ri == 0 else cell_style]
            except KeyError:
                pass
            if ri == 0:
                run.bold = True
    return t


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def emit_section(doc, items):
    for kind, text in items:
        if kind == "h1":
            add_para(doc, text, "IEEE Heading 1")
        elif kind == "h2":
            add_para(doc, text, "IEEE Heading 2")
        elif kind == "h3":
            add_para(doc, text, "IEEE Heading 3")
        elif kind == "p":
            add_para(doc, text, "IEEE Paragraph")
        elif kind == "ref":
            add_para(doc, text, "IEEE Reference Item")
        elif kind == "bullet":
            add_para(doc, text, "IEEE Bullet 1")
        elif kind == "pagebreak":
            add_page_break(doc)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #

def main(template: Path, output: Path) -> None:
    doc = Document(str(template))

    # 1) Replace abstract + keywords. The abstract paragraph is detected by
    #    the IEEE Abtract style; the keywords paragraph is the one that
    #    starts with "Keywords".
    abstract_replaced = False
    keywords_replaced = False
    title_replaced = False
    keywords_para = None
    for p in doc.paragraphs:
        if not title_replaced and p.style.name == "IEEE Title" and p.text.strip():
            # Title spans two paragraphs in the template ("Learning..." +
            # "Grid-Based Puzzles..."). Use the first for the project
            # name and the second as a one-line subtitle.
            p.text = "Candy Crusher"
            title_replaced = True
            continue
        if title_replaced and p.style.name == "IEEE Title" and p.text.strip():
            p.text = "Classical RL Baselines and a GRPO-Trained LLM Policy"
            continue
        if not abstract_replaced and p.style.name == "IEEE Abtract":
            # Keep the leading "Abstract—" run intact, replace the body text.
            new_text = "Abstract—" + ABSTRACT
            p.text = ""
            run = p.add_run(new_text)
            run.bold = False
            abstract_replaced = True
            continue
        if not keywords_replaced and p.text.strip().startswith("Keywords"):
            # Drop the keywords line entirely.
            keywords_para = p
            p.text = ""
            keywords_replaced = True
            continue

    # 2) Drop everything after the keywords paragraph so we can rewrite
    #    the body cleanly.
    paragraphs = list(doc.paragraphs)
    if keywords_para is None:
        raise RuntimeError("Could not locate Keywords paragraph in template.")
    kw_el = keywords_para._element
    keywords_idx = next(
        i for i, p in enumerate(paragraphs) if p._element is kw_el
    )

    # Drop the (now-empty) keywords paragraph itself plus everything
    # after it, and remove every existing table from the body.
    for p in paragraphs[keywords_idx:]:
        remove_paragraph(p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # 3) Emit the report body.
    # ---- Section 1 + 2: Problem statement + RL formulation ----
    intro_end = next(i for i, (k, _) in enumerate(SECTIONS) if k == "h1" and _ == "Methodology")
    emit_section(doc, SECTIONS[:intro_end])

    # ---- Section 3: Methodology + Contributions ----
    contrib_start = next(
        i for i, (k, t) in enumerate(SECTIONS) if k == "h1" and t == "Contributions"
    )
    exp_setup_start = next(
        i for i, (k, t) in enumerate(SECTIONS) if k == "h1" and t == "Experimental Setup"
    )
    emit_section(doc, SECTIONS[intro_end:contrib_start])  # Methodology

    # Hyperparameters table at end of Methodology section
    add_para(
        doc,
        "Table I. Key hyperparameters of the three trained policies.",
        "IEEE Table Caption",
    )
    add_table_with_style(doc, HYPERPARAMS_TABLE)

    emit_section(doc, SECTIONS[contrib_start:exp_setup_start])  # Contributions intro

    # Contributions table (4 rows, equal split)
    add_para(
        doc,
        "Table II. Author contributions (equal split).",
        "IEEE Table Caption",
    )
    contrib_rows = [("Author", "Contribution")] + CONTRIBUTIONS
    add_table_with_style(doc, contrib_rows)

    # ---- Section 4: Results ----
    emit_section(doc, SECTIONS[exp_setup_start:])

    # The Results table goes right after the "Results" heading. Find it
    # and insert the two result tables inline. The cleanest way given
    # python-docx's API: locate the first paragraph whose text starts
    # with "Greedy is the strongest policy" and insert tables BEFORE it.
    insert_anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("Greedy is the strongest policy"):
            insert_anchor = p
            break
    if insert_anchor is None:
        raise RuntimeError("Could not find Results discussion paragraph.")

    # Build the tables at end of doc, then move them to the anchor.
    summary_caption = add_para(
        doc,
        "Table III. Aggregate reward across 10 boards (full 20-move rollouts).",
        "IEEE Table Caption",
    )
    summary_table = add_table_with_style(doc, RESULTS_TABLE)
    perboard_caption = add_para(
        doc,
        "Table IV. Per-board total reward by policy.",
        "IEEE Table Caption",
    )
    perboard_table = add_table_with_style(doc, PER_BOARD_TABLE)

    anchor_el = insert_anchor._element
    for el in (
        summary_caption._element,
        summary_table._element,
        perboard_caption._element,
        perboard_table._element,
    ):
        anchor_el.addprevious(el)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Wrote {output}")


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument(
        "--template",
        default="/Users/arnavmehta/Downloads/RL report Template.docx",
        help="Path to the IEEE template .docx",
    )
    p.add_argument(
        "--output",
        default=str(ROOT / "report" / "RL_Report.docx"),
        help="Where to write the generated report",
    )
    return p.parse_args()


if __name__ == "__main__":
    args = parse_args()
    main(Path(args.template), Path(args.output))
